"""
論文前置頁面 Word 轉換腳本

使用方式：
1. 安裝 python-docx: pip install python-docx
2. 執行腳本: python convert_to_word.py

輸出：
- 00a_謝辭.docx
- 00b_中文摘要.docx
- 00b_英文摘要.docx
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 設定輸出目錄
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
CHAPTERS_DIR = os.path.join(PROJECT_DIR, "docs", "00_thesis", "chapters")
OUTPUT_DIR = os.path.join(PROJECT_DIR, "docs", "00_thesis", "word")

# 確保輸出目錄存在
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_styled_document():
    """建立已設定好格式的空白文件"""
    doc = Document()

    # 設定頁面邊界：上3、下2、左3、右3 cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # 設定預設字型
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 設定中文字型（標楷體）
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    # 設定行距 1.5 倍
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)

    return doc


def add_title(doc, title_text, is_centered=True):
    """新增標題"""
    title = doc.add_paragraph()
    run = title.add_run(title_text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    if is_centered:
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 標題後空一行
    doc.add_paragraph()

    return title


def add_paragraph(doc, text, first_line_indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """新增段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    para.alignment = alignment

    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.85)  # 約 2 字元

    return para


def add_keywords(doc, keywords_label, keywords_text):
    """新增關鍵詞"""
    para = doc.add_paragraph()

    # 關鍵詞標籤（粗體）
    run_label = para.add_run(keywords_label)
    run_label.font.size = Pt(12)
    run_label.font.bold = True
    run_label.font.name = 'Times New Roman'
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    # 關鍵詞內容
    run_text = para.add_run(keywords_text)
    run_text.font.size = Pt(12)
    run_text.font.name = 'Times New Roman'
    run_text._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    return para


def create_acknowledgment():
    """建立謝辭文件"""
    doc = create_styled_document()

    add_title(doc, "謝辭")

    paragraphs = [
        "脫離學業十年有餘，回想當初考上心目中的臺北科技大學時，抱著由你玩四年的心態，並沒有將重心放在課業上。然而，大學時期工程數學課堂上，楊士萱教授曾說過一句話：「不要對自己太好」，這句話一直銘記在心，卻是出了社會、經歷了職場的磨練後，才真正深刻體悟其中的意義。十多年後，帶著截然不同的心態重返學術殿堂，這一次，我格外珍惜每一堂課、每一次學習的機會。",
        "首先，誠摯感謝指導教授許揚老師的悉心指導。當初選擇軟體工程作為研究方向，正是因為這個領域與我十多年的軟體開發職涯高度契合，這是一個不會後悔的選擇。教授深知在職學生的處境，鼓勵我從工作實務中尋找研究題目，讓學術研究與職場經驗得以相互印證，這樣的指導方式讓我受益良多。",
        "感謝公司大安聯合醫事檢驗所提供的在職進修福利，這份制度成為我重返學術領域的重要推力。公司全額補助學費，大幅降低了經濟上的顧慮，使我能夠專注於學業。同時，順利取得碩士學位後的額外加薪制度，更體現了公司對員工自我提升的重視。我始終相信，所學到的知識是別人帶不走的，而公司投資員工成長、員工回饋所學於工作，正是一個雙贏的局面。",
        "最後，我要特別感謝我的妻子。她是一位獨立且有能力的人，在我每天早出晚歸的求學期間，默默承擔了許多家庭的責任與付出。更令人欣慰的是，在我就讀碩士的第一年，她也報名了國外的碩士進修課程，大兒子進入國小一年級，小兒子進入幼稚園——一家四口同時都是學生，整個家庭充滿了學習的氛圍。這份共同成長的經歷，是這段求學旅程中最珍貴的收穫。",
        "謹以此論文，獻給所有支持我的人。",
    ]

    for para_text in paragraphs:
        add_paragraph(doc, para_text)

    # 儲存
    output_path = os.path.join(OUTPUT_DIR, "00a_謝辭.docx")
    doc.save(output_path)
    print(f"已建立：{output_path}")
    return output_path


def create_chinese_abstract():
    """建立中文摘要文件"""
    doc = create_styled_document()

    add_title(doc, "摘要")

    # 論文題目
    title_para = doc.add_paragraph()
    run_label = title_para.add_run("論文題目：")
    run_label.font.bold = True
    run_label.font.size = Pt(12)
    run_label.font.name = 'Times New Roman'
    run_label._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    run_title = title_para.add_run("縱向健檢資料與變化量特徵之三高疾病風險預測：多模型比較研究")
    run_title.font.size = Pt(12)
    run_title.font.name = 'Times New Roman'
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    doc.add_paragraph()  # 空行

    paragraphs = [
        "三高疾病（高血壓、高血糖、高血脂）是全球主要的慢性疾病，也是心血管疾病的關鍵可控風險因子。在台灣，40 歲以上國人三高盛行率介於 16.4%–38.3%，且約有 4 至 7 成患者不知道自身已罹病，凸顯早期預測的重要性。然而，現有風險評估方法多仰賴單一時間點的檢驗數據，未能充分利用縱向健檢資料中蘊含的動態資訊。",
        "本研究使用公開於 Dryad 資料庫的縱向健檢資料集（Luo et al., 2024），涵蓋中國杭州市 6,056 位 40 歲以上社區成人，追蹤期間為 2010 至 2018 年。研究採用三時間點縱貫設計（Y-2、Y-1、Y0，間隔約兩年），以 Y-2 與 Y-1 的健檢指標及其變化量（Δ 特徵）作為輸入，預測 Y0 時的三高疾病狀態。透過滑動窗口法，共產生 13,514 筆建模紀錄。本研究系統性比較八種模型：傳統統計方法（Logistic Regression、Naive Bayes、LDA）、樹模型（Decision Tree、Random Forest、XGBoost）、核方法（SVM）及神經網路（MLP），並另以符號回歸（PySR）探討可解釋性。實驗採用 StratifiedGroupKFold 五折交叉驗證，確保同一受檢者不同時出現於訓練集與測試集，同時維持各折的類別比例。",
        "主要研究發現如下：（1）Logistic Regression 在三項預測任務中皆展現穩定且優異的表現，高血糖預測 AUC 達 0.938；（2）Δ 特徵在資料受限情境下（僅有 Y-1 資料）可帶來 1.5%–2.3% 的 AUC 提升，且在各疾病 Top 10 重要特徵中佔比達 30–50%；（3）SHAP 分析揭示疾病特異性預測因子，如高血壓以 SBP/DBP 為主、高血糖以 FBG 為主、高血脂以 TC 與 ΔeGFR 最為重要；（4）符號回歸發現極簡公式 0.114 × FBG_Y-1 即可達到 AUC 0.943，展現可解釋 AI 的實務潛力；（5）健檢次數實驗顯示累積更多健檢紀錄有助於提升預測性能，三項疾病皆呈正相關趨勢；（6）僅使用 Top 5 特徵，AUC 降幅皆小於 0.5%，證實精簡模型的可行性。",
        "本研究的核心貢獻在於：以容易取得的縱向健檢資料，透過簡單的特徵工程與線性模型，即可達到臨床可用的預測性能，無需複雜的深度學習模型或昂貴的檢驗項目，適合在基層醫療單位實施早期預警系統。",
    ]

    for para_text in paragraphs:
        add_paragraph(doc, para_text)

    doc.add_paragraph()  # 空行

    # 關鍵詞
    add_keywords(doc, "關鍵詞：", "三高疾病、高血壓、高血糖、高血脂、機器學習、縱向資料、變化量特徵、SHAP 可解釋性、符號回歸")

    # 儲存
    output_path = os.path.join(OUTPUT_DIR, "00b_中文摘要.docx")
    doc.save(output_path)
    print(f"已建立：{output_path}")
    return output_path


def create_english_abstract():
    """建立英文摘要文件"""
    doc = create_styled_document()

    add_title(doc, "Abstract")

    # 論文題目
    title_para = doc.add_paragraph()
    run_label = title_para.add_run("Title: ")
    run_label.font.bold = True
    run_label.font.size = Pt(12)
    run_label.font.name = 'Times New Roman'

    run_title = title_para.add_run("Predicting Hypertension, Hyperglycemia, and Dyslipidemia Using Longitudinal Health Checkup Data with Delta Features: A Multi-Model Comparative Study")
    run_title.font.size = Pt(12)
    run_title.font.name = 'Times New Roman'

    doc.add_paragraph()  # 空行

    paragraphs = [
        "Hypertension, hyperglycemia, and dyslipidemia—collectively known as the \"three highs\"—are major chronic diseases and key modifiable risk factors for cardiovascular disease. Conventional risk assessment methods rely on single-timepoint examination data, failing to capture the dynamic health trajectories embedded in longitudinal records. This study aims to leverage longitudinal health checkup data and delta features (Δ features) to predict future disease risk while balancing predictive accuracy and model interpretability.",
        "We utilized a publicly available longitudinal health checkup dataset (Luo et al., 2024, Dryad repository) comprising 6,056 community-dwelling adults aged 40 and above from Hangzhou, China, followed between 2010 and 2018. A three-timepoint longitudinal design (Y-2, Y-1, Y0; approximately two-year intervals) was adopted, using health indicators and their changes (Δ features) from Y-2 and Y-1 as input to predict disease status at Y0. A sliding window approach generated 13,514 modeling records. Eight models were systematically compared: traditional statistical methods (Logistic Regression, Naive Bayes, LDA), tree-based models (Decision Tree, Random Forest, XGBoost), kernel method (SVM), and neural network (MLP), with symbolic regression (PySR) additionally employed to explore interpretability. StratifiedGroupKFold 5-fold cross-validation ensured no data leakage across participants while maintaining class proportions in each fold.",
        "Key findings include: (1) Logistic Regression demonstrated consistently strong performance across all three tasks, achieving AUC 0.938 for hyperglycemia prediction; (2) Δ features provided 1.5%–2.3% AUC improvement in data-limited scenarios (Y-1 only) and comprised 30–50% of the top 10 important features; (3) SHAP analysis revealed disease-specific predictors—SBP/DBP for hypertension, FBG for hyperglycemia, and TC/ΔeGFR for dyslipidemia; (4) symbolic regression discovered a minimal formula (0.114 × FBG_Y-1) achieving AUC 0.943, demonstrating the practical potential of interpretable AI; (5) increasing health checkup frequency was positively correlated with prediction performance across all three diseases; and (6) using only the top 5 features resulted in less than 0.5% AUC decrease, confirming the feasibility of parsimonious models.",
        "This study demonstrates that readily accessible longitudinal health checkup data, combined with simple feature engineering and linear models, can achieve clinically useful predictive performance without complex deep learning models or expensive laboratory tests, making it suitable for deployment as an early warning system in primary healthcare settings.",
    ]

    for para_text in paragraphs:
        add_paragraph(doc, para_text)

    doc.add_paragraph()  # 空行

    # 關鍵詞
    add_keywords(doc, "Keywords: ", "Hypertension, Hyperglycemia, Dyslipidemia, Machine Learning, Longitudinal Data, Delta Features, SHAP Interpretability, Symbolic Regression")

    # 儲存
    output_path = os.path.join(OUTPUT_DIR, "00b_英文摘要.docx")
    doc.save(output_path)
    print(f"已建立：{output_path}")
    return output_path


def main():
    print("=" * 50)
    print("論文前置頁面 Word 轉換工具")
    print("=" * 50)
    print(f"\n輸出目錄：{OUTPUT_DIR}\n")

    # 建立各文件
    create_acknowledgment()
    create_chinese_abstract()
    create_english_abstract()

    print("\n" + "=" * 50)
    print("轉換完成！")
    print("=" * 50)
    print("\n請在 Word 中開啟檔案，進行以下檢查：")
    print("1. 字型：中文標楷體 12pt、英文 Times New Roman 12pt")
    print("2. 行距：1.5 倍行高")
    print("3. 邊界：上3、下2、左3、右3 cm")
    print("4. 首行縮排：約 2 字元")


if __name__ == "__main__":
    main()
